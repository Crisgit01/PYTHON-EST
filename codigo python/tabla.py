from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

def agregar_tabla_a_docx(doc, numero_tabla, descripcion_tabla, df):
    # Crear un título para la TABLA
    paragraph = doc.add_paragraph()
    paragraph.add_run(f'Tabla {numero_tabla}. ').bold = True
    paragraph.add_run(descripcion_tabla)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Centrar el título

    # Agregar una tabla al documento con 1 fila para los encabezados y n filas para los datos
    table = doc.add_table(rows=df.shape[0] + 1, cols=len(df.columns))  # +1 para incluir la fila de encabezados

    # Establecer estilos para la tabla y las celdas
    table.style = 'Table Grid'  # Aplicar un estilo de tabla con líneas

    # Agregar los encabezados de columna y centrar el texto
    for i, column_name in enumerate(df.columns):
        cell = table.cell(0, i)
        cell.text = column_name
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Centrar el texto

    # Agregar los datos del DataFrame a la tabla y centrar el texto
    for row in range(df.shape[0]):
        for col in range(df.shape[1]):
            cell = table.cell(row + 1, col)
            cell.text = str(df.iloc[row, col])
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Centrar el texto
