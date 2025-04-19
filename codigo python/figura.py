from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def agregar_figura_a_docx_existente(doc, ubicacion_imagen, numero_figura, descripcion_figura, ancho, largo):
    # Agregar un párrafo vacío para crear un espacio entre el título y el subtítulo
    doc.add_paragraph('')

    # Agregar imagen a WORD
    doc.add_picture(ubicacion_imagen, width=Inches(ancho), height=Inches(largo))

    # Centrar imagen
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Crear un título para la imagen
    paragraph = doc.add_paragraph()
    paragraph.add_run(f'Figura {numero_figura}. ').bold = True
    paragraph.add_run(descripcion_figura)