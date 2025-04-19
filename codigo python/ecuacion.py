import matplotlib.pyplot as plt
from io import BytesIO
import base64
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def agregar_ecuacion_latex_a_docx(doc, latex_equation):
    # Función para convertir una ecuación LaTeX en una imagen PNG
    def latex_to_image(latex):
        fig, ax = plt.subplots(figsize=(7, 0.55))
        ax.text(0.4, 0.4, f'${latex}$', size=10, ha='center', va='center')
        ax.axis('off')
        buffer = BytesIO()
        plt.savefig(buffer, format='png', bbox_inches='tight', pad_inches=0.0, dpi=300)
        buffer.seek(0)
        image_data = base64.b64encode(buffer.read()).decode()
        plt.close()
        return image_data

    # Convertir la ecuación LaTeX en una imagen
    equation_image_data = latex_to_image(latex_equation)

    # Insertar la imagen en el documento
    doc.add_paragraph().add_run().add_picture(BytesIO(base64.b64decode(equation_image_data)))

    # Centrar imagen
    last_paragraph = doc.paragraphs[-1] 
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER